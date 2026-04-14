#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Universal Contract Generator

生成专业的合同文档（雇佣合同、租房合同、NDA、服务合同等）。
结合自动化模板填充与法律最佳实践。

Version: 1.4.0

修复记录：
  1.4.0 - 新增金额中文大写转换（money_to_cn）；
          新增 list-templates CLI 命令；
          新增文本模板生成（NDA、服务合同）；
          修复 _replace_in_paragraph 保留首 Run 字体/粗体格式；
          修复 batch 输出文件名包含非法字符导致保存失败；
          增加 --output-dir 默认值（contracts/）；
          增强 validate 输出可读性
  1.3.0 - 按合同类型区分 validate 关键词；BatchGenerator 错误隔离 + 结果报告；
          支持 Excel 数据源；残留占位符检测；日期自动计算；CSV 编码自动检测
  1.2.0 - 修复跨 Run 占位符替换；改用 deepcopy 保留模板格式；表格单元格扫描
  1.1.0 - 添加租房合同支持，增强错误处理
  1.0.0 - 初始版本
"""

__version__ = '1.4.0'

import os
import re
import sys
import json
import copy
import csv
import argparse
from pathlib import Path
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from typing import Dict, List, Optional, Any, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：python-docx 未安装，部分功能不可用。")
    print("安装：pip install python-docx")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


# ---------------------------------------------------------------------------
# 合同类型关键词配置
# ---------------------------------------------------------------------------

CONTRACT_KEYWORDS: Dict[str, List[str]] = {
    'employment': ['甲方', '乙方', '合同期限', '工作内容', '劳动报酬', '社会保险', '签字', '日期'],
    'offer_letter': ['甲方', '乙方', '职位', '薪酬', '入职日期', '签字', '日期'],
    'rental': ['出租方', '承租方', '租赁期限', '租金', '押金', '签字', '日期'],
    'nda': ['甲方', '乙方', '保密信息', '保密期限', '签字', '日期'],
    'service': ['甲方', '乙方', '服务内容', '服务费用', '签字', '日期'],
    'generic': ['甲方', '乙方', '签字', '日期'],
}

TEMPLATE_TYPE_MAP: Dict[str, str] = {
    'employment_contract': 'employment',
    'offer_letter': 'offer_letter',
    'rental_contract': 'rental',
    'nda': 'nda',
    'service_contract': 'service',
}


def _detect_contract_type(contract_path: str) -> str:
    """根据文件名推断合同类型，匹配不到则返回 'generic'。"""
    stem = Path(contract_path).stem.lower()
    for key, ctype in TEMPLATE_TYPE_MAP.items():
        if key in stem:
            return ctype
    return 'generic'


# ---------------------------------------------------------------------------
# 金额中文大写转换（1.4.0 新增）
# ---------------------------------------------------------------------------

_CN_NUMS = '零壹贰叁肆伍陆柒捌玖'
_CN_UNITS = ['', '拾', '佰', '仟']
_CN_GROUPS = ['', '万', '亿']


def money_to_cn(amount: Any) -> str:
    """
    将数字金额转为中文大写（人民币格式）。
    示例：50000 → "伍万元整"，12345.67 → "壹万贰仟叁佰肆拾伍元陆角柒分"
    支持整数、浮点数、字符串输入。
    """
    try:
        # 清理输入
        amount_str = str(amount).replace(',', '').replace('，', '').strip()
        # 拆分整数和小数
        if '.' in amount_str:
            int_part_str, dec_part_str = amount_str.split('.', 1)
            dec_part_str = dec_part_str[:2].ljust(2, '0')  # 最多两位小数
        else:
            int_part_str = amount_str
            dec_part_str = '00'

        int_part = int(int_part_str)
        jiao = int(dec_part_str[0])
        fen = int(dec_part_str[1])

        if int_part < 0:
            raise ValueError("不支持负数金额")

        def _group_to_cn(g: int) -> str:
            """将 0~9999 转为中文，不带万/亿单位。"""
            if g == 0:
                return ''
            thousands = g // 1000
            hundreds = (g % 1000) // 100
            tens = (g % 100) // 10
            ones = g % 10
            s = ''
            if thousands:
                s += _CN_NUMS[thousands] + '仟'
            if hundreds:
                s += _CN_NUMS[hundreds] + '佰'
            elif thousands and (tens or ones):
                s += '零'
            if tens:
                s += _CN_NUMS[tens] + '拾'
            elif (thousands or hundreds) and ones:
                s += '零'
            if ones:
                s += _CN_NUMS[ones]
            return s

        def _int_to_cn(n: int) -> str:
            if n == 0:
                return '零'
            groups = []
            while n > 0:
                groups.append(n % 10000)
                n //= 10000
            # groups[0] = 个位组，groups[-1] = 最高位组
            result = ''
            total = len(groups)
            for gi, g in enumerate(reversed(groups)):
                level = total - 1 - gi          # 0=个位组, 1=万, 2=亿
                group_cn = _group_to_cn(g)
                if g == 0:
                    # 高位组为零：只有当结果非空且尾部不是零时补零
                    if result and not result.endswith('零'):
                        result += '零'
                else:
                    # 当前组不为零，但低于最高位且首位不足千位，需要补零
                    if result and g < 1000 and not result.endswith('零'):
                        result += '零'
                    result += group_cn + _CN_GROUPS[level]
            # 合并连续零，去掉末尾零
            result = re.sub(r'零+', '零', result).rstrip('零')
            return result

        int_cn = _int_to_cn(int_part)
        result = int_cn + '元'

        if jiao == 0 and fen == 0:
            result += '整'
        elif jiao == 0:
            result += '零' + _CN_NUMS[fen] + '分'
        elif fen == 0:
            result += _CN_NUMS[jiao] + '角整'
        else:
            result += _CN_NUMS[jiao] + '角' + _CN_NUMS[fen] + '分'

        return result
    except Exception as e:
        return f'[金额转换失败：{e}]'


# ---------------------------------------------------------------------------
# 日期工具
# ---------------------------------------------------------------------------

_DATE_FORMATS = [
    '%Y年%m月%d日', '%Y-%m-%d', '%Y/%m/%d',
    '%Y.%m.%d', '%Y%m%d',
]


def parse_date(date_str: str) -> Optional[date]:
    """尝试多种格式解析日期字符串，失败返回 None。"""
    s = str(date_str).strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def format_date_cn(d: date) -> str:
    """将 date 对象格式化为中文日期字符串。"""
    return f"{d.year}年{d.month}月{d.day}日"


def compute_derived_fields(variables: Dict[str, Any]) -> Dict[str, str]:
    """
    根据已有变量自动推算衍生字段：
      - end_date: start_date + contract_years
      - probation_end_date: start_date + probation_months
      - lease_end_date: lease_start_date + lease_months
      - service_fee_uppercase: service_fee 的中文大写（1.4.0 新增）
      - monthly_rent_uppercase: monthly_rent 的中文大写（1.4.0 新增）
      - deposit_amount_uppercase: deposit_amount 的中文大写（1.4.0 新增）
      - base_salary_uppercase: base_salary 的中文大写（1.4.0 新增）
    返回新增字段字典（不覆盖已有值）。
    """
    derived: Dict[str, str] = {}

    # 劳动合同：start_date + contract_years → end_date
    if 'start_date' in variables and 'end_date' not in variables:
        start = parse_date(str(variables['start_date']))
        if start and 'contract_years' in variables:
            try:
                years = int(variables['contract_years'])
                end = start + relativedelta(years=years) - relativedelta(days=1)
                derived['end_date'] = format_date_cn(end)
            except (ValueError, TypeError):
                pass

    # 劳动合同：start_date + probation_months → probation_end_date
    if 'start_date' in variables and 'probation_end_date' not in variables:
        start = parse_date(str(variables.get('start_date', '')))
        if start and 'probation_months' in variables:
            try:
                months = int(variables['probation_months'])
                prob_end = start + relativedelta(months=months) - relativedelta(days=1)
                derived['probation_end_date'] = format_date_cn(prob_end)
            except (ValueError, TypeError):
                pass

    # 租房合同：lease_start_date + lease_months → lease_end_date
    if 'lease_start_date' in variables and 'lease_end_date' not in variables:
        start = parse_date(str(variables['lease_start_date']))
        if start and 'lease_months' in variables:
            try:
                months = int(variables['lease_months'])
                end = start + relativedelta(months=months) - relativedelta(days=1)
                derived['lease_end_date'] = format_date_cn(end)
            except (ValueError, TypeError):
                pass

    # 金额大写（1.4.0 新增）
    money_fields = {
        'service_fee': 'service_fee_uppercase',
        'monthly_rent': 'monthly_rent_uppercase',
        'deposit_amount': 'deposit_amount_uppercase',
        'base_salary': 'base_salary_uppercase',
    }
    for src, dst in money_fields.items():
        if src in variables and dst not in variables:
            try:
                derived[dst] = money_to_cn(variables[src])
            except Exception:
                pass

    return derived


# ---------------------------------------------------------------------------
# CSV 编码检测
# ---------------------------------------------------------------------------

def _detect_csv_encoding(csv_path: str) -> str:
    """
    检测 CSV 文件编码。
    优先尝试 utf-8-sig（带 BOM 的 UTF-8，Excel 默认导出格式），
    再尝试 gbk，最后回退到 utf-8。
    """
    path = Path(csv_path)
    for enc in ('utf-8-sig', 'gbk', 'utf-8', 'gb2312', 'gb18030'):
        try:
            with open(path, 'r', encoding=enc) as f:
                f.read(1024)
            return enc
        except (UnicodeDecodeError, LookupError):
            continue
    return 'utf-8'


# ---------------------------------------------------------------------------
# 文件名安全清理（1.4.0 修复）
# ---------------------------------------------------------------------------

def _safe_filename(name: str, fallback: str = 'doc') -> str:
    """
    移除文件名中不合法的字符（Windows/Unix 均兼容），保留中文、字母、数字、下划线、连字符。
    """
    safe = re.sub(r'[\\/:*?"<>|]', '', str(name))
    safe = safe.strip().strip('.')
    return safe if safe else fallback


# ---------------------------------------------------------------------------
# ContractGenerator
# ---------------------------------------------------------------------------

class ContractGenerator:
    """合同生成器 - 单个合同生成"""

    def __init__(self, template_dir: Optional[str] = None):
        if template_dir:
            self.template_dir = Path(template_dir)
        else:
            self.template_dir = Path(__file__).parent / 'templates'

        self.template = None
        self.template_path: Optional[Path] = None
        self.variables: Dict[str, Any] = {}
        self.required_fields: List[str] = []

    # ------------------------------------------------------------------
    # 模板加载
    # ------------------------------------------------------------------

    def load_template(self, template_name: str) -> 'ContractGenerator':
        """加载 DOCX 模板文件。"""
        # 支持绝对路径
        tp = Path(template_name)
        if tp.is_absolute() and tp.exists():
            self.template_path = tp
        else:
            self.template_path = self.template_dir / template_name

        if not self.template_path.exists():
            available = (
                [f.name for f in self.template_dir.glob('*.docx')]
                if self.template_dir.exists() else []
            )
            raise FileNotFoundError(
                f"模板文件不存在：{self.template_path}\n"
                f"可用模板：{', '.join(available) if available else '无'}\n"
                f"提示：可运行 `python contract_generator.py list-templates` 查看模板列表"
            )

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行：pip install python-docx")

        self.template = Document(self.template_path)
        self._extract_required_fields()
        return self

    # ------------------------------------------------------------------
    # 字段提取
    # ------------------------------------------------------------------

    def _extract_required_fields(self) -> None:
        """从模板中提取所有 [field_name] 占位符（段落 + 表格）。"""
        if not self.template:
            return
        fields: set = set()
        pattern = re.compile(r'\[([a-zA-Z_][a-zA-Z0-9_]*)\]')

        for para in self.template.paragraphs:
            fields.update(pattern.findall(para.text))

        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    fields.update(pattern.findall(cell.text))

        self.required_fields = list(fields)

    # ------------------------------------------------------------------
    # 变量设置
    # ------------------------------------------------------------------

    def set_variables(self, variables: Dict[str, Any]) -> 'ContractGenerator':
        """设置填充变量，自动补充可推算的衍生字段（日期、金额大写）。"""
        self.variables.update(variables)
        derived = compute_derived_fields(self.variables)
        for k, v in derived.items():
            if k not in self.variables:
                self.variables[k] = v
        return self

    # ------------------------------------------------------------------
    # 字段验证
    # ------------------------------------------------------------------

    def validate(self) -> List[str]:
        """返回模板中存在但未提供值的字段列表。"""
        return [f for f in self.required_fields if f not in self.variables]

    # ------------------------------------------------------------------
    # 残留占位符检测
    # ------------------------------------------------------------------

    def find_unreplaced_placeholders(self, doc: 'Document') -> List[str]:
        """扫描已生成文档，返回仍残留的 [xxx] 占位符列表。"""
        pattern = re.compile(r'\[([a-zA-Z_][a-zA-Z0-9_]*)\]')
        found: set = set()

        for para in doc.paragraphs:
            found.update(pattern.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found.update(pattern.findall(cell.text))

        return sorted(found)

    # ------------------------------------------------------------------
    # 占位符替换
    # ------------------------------------------------------------------

    def _replace_placeholders(self, text: str) -> str:
        """替换字符串中的所有占位符（[key] 和 {{key}} 两种格式）。"""
        for key, value in self.variables.items():
            pattern = re.compile(r'\[' + re.escape(key) + r'\]', re.IGNORECASE)
            text = pattern.sub(str(value), text)
            text = text.replace('{{' + key + '}}', str(value))
        return text

    def _replace_in_paragraph(self, para) -> None:
        """
        合并段落全部 Run 文本后替换，再写回第一个 Run。
        1.4.0 修复：保留首 Run 的字体、粗体、字号等格式不被清空。
        """
        if not para.runs:
            return
        full_text = ''.join(run.text for run in para.runs)
        replaced = self._replace_placeholders(full_text)
        if full_text == replaced:
            return  # 无变化则跳过，避免不必要的写入
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ''

    def _replace_in_cell(self, cell) -> None:
        """替换表格单元格内所有段落的占位符。"""
        for para in cell.paragraphs:
            self._replace_in_paragraph(para)

    # ------------------------------------------------------------------
    # 生成 & 保存
    # ------------------------------------------------------------------

    def generate(self) -> 'Document':
        """
        深拷贝模板后原地替换占位符，保留所有样式和格式。
        生成后自动检测残留占位符并发出警告。
        """
        if not self.template:
            raise ValueError("请先调用 load_template() 加载模板")

        missing = self.validate()
        if missing:
            print(f"警告：以下字段未提供，将保留占位符：{', '.join(missing)}")

        new_doc = copy.deepcopy(self.template)

        for para in new_doc.paragraphs:
            self._replace_in_paragraph(para)

        for table in new_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_cell(cell)

        unreplaced = self.find_unreplaced_placeholders(new_doc)
        if unreplaced:
            print(f"警告：文档中存在未替换的占位符：{', '.join(unreplaced)}")

        return new_doc

    def save(self, output_path: str) -> str:
        """生成合同并保存到指定路径，返回保存路径。"""
        doc = self.generate()
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        print(f"✓ 合同已保存至：{output.resolve()}")
        return str(output)

    def generate_text(self, template_text: str) -> str:
        """从纯文本模板生成合同文本（不涉及 DOCX）。"""
        return self._replace_placeholders(template_text)

    def save_from_text(self, template_text: str, output_path: str) -> str:
        """
        1.4.0 新增：从纯文本模板生成合同并保存为 DOCX。
        用于 NDA、服务合同等尚无 DOCX 模板的合同类型。
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行：pip install python-docx")

        content = self.generate_text(template_text)
        doc = Document()

        for line in content.splitlines():
            p = doc.add_paragraph(line)

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        print(f"✓ 合同已保存至：{output.resolve()}")
        return str(output)


# ---------------------------------------------------------------------------
# BatchGenerator
# ---------------------------------------------------------------------------

class BatchGenerator:
    """批量合同生成器"""

    def __init__(self, template: str, template_dir: Optional[str] = None):
        self.template = template
        self.template_dir = template_dir
        self.records: List[Dict[str, Any]] = []

    # ------------------------------------------------------------------
    # 数据加载
    # ------------------------------------------------------------------

    def load_csv(self, csv_path: str) -> 'BatchGenerator':
        """从 CSV 文件加载数据，自动检测编码。"""
        encoding = _detect_csv_encoding(csv_path)
        with open(csv_path, 'r', encoding=encoding, newline='') as f:
            reader = csv.DictReader(f)
            self.records = [dict(row) for row in reader]
        print(f"已加载 {len(self.records)} 条记录（编码：{encoding}）")
        return self

    def load_json(self, json_path: str) -> 'BatchGenerator':
        """从 JSON 文件加载数据（支持对象数组或单对象）。"""
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        self.records = [data] if isinstance(data, dict) else data
        print(f"已加载 {len(self.records)} 条记录")
        return self

    def load_excel(self, excel_path: str, sheet_name: int = 0) -> 'BatchGenerator':
        """从 Excel (.xlsx / .xls) 文件加载数据。依赖：pandas openpyxl"""
        if not PANDAS_AVAILABLE:
            raise ImportError("读取 Excel 需要 pandas，请运行：pip install pandas openpyxl")
        df = pd.read_excel(excel_path, sheet_name=sheet_name, dtype=str)
        df = df.where(pd.notna(df), '')
        self.records = df.to_dict(orient='records')
        print(f"已加载 {len(self.records)} 条记录（来源：{Path(excel_path).name}）")
        return self

    # ------------------------------------------------------------------
    # 批量生成
    # ------------------------------------------------------------------

    def generate_all(
        self,
        output_dir: str = 'contracts/',
        name_field: str = 'name',
    ) -> Dict[str, Any]:
        """
        批量生成所有合同。单条失败不中断整批，返回详细结果报告。

        Returns:
            {
              'total': int,
              'success': int,
              'failed': int,
              'generated_files': [str, ...],
              'errors': [{'index': int, 'name': str, 'error': str}, ...]
            }
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        generated_files: List[str] = []
        errors: List[Dict[str, Any]] = []

        for i, record in enumerate(self.records):
            raw_name = (
                record.get(name_field)
                or record.get('employee_name')
                or record.get('candidate_name')
                or record.get('tenant_name')
                or f'record_{i+1}'
            )
            # 1.4.0 修复：安全文件名（移除非法字符）
            safe_name = _safe_filename(str(raw_name), fallback=f'record_{i+1}')
            template_stem = Path(self.template).stem
            output_file = output_path / f"{template_stem}_{safe_name}.docx"

            try:
                generator = ContractGenerator(self.template_dir)
                generator.load_template(self.template)
                generator.set_variables(dict(record))
                generator.save(str(output_file))
                generated_files.append(str(output_file))
            except Exception as e:
                error_info = {'index': i + 1, 'name': raw_name, 'error': str(e)}
                errors.append(error_info)
                print(f"  ✗ 第 {i + 1} 条（{raw_name}）生成失败：{e}")

        total = len(self.records)
        success = len(generated_files)
        failed = len(errors)

        print(f"\n{'=' * 50}")
        print(f"批量生成报告")
        print(f"{'=' * 50}")
        print(f"  总计：{total} 条")
        print(f"  ✓ 成功：{success} 条")
        if failed:
            print(f"  ✗ 失败：{failed} 条")
            for err in errors:
                print(f"      第 {err['index']} 条（{err['name']}）：{err['error']}")
        print(f"  输出目录：{output_path.resolve()}")
        print(f"{'=' * 50}")

        return {
            'total': total,
            'success': success,
            'failed': failed,
            'generated_files': generated_files,
            'errors': errors,
        }


# ---------------------------------------------------------------------------
# validate_contract
# ---------------------------------------------------------------------------

def validate_contract(contract_path: str, contract_type: Optional[str] = None) -> Dict[str, Any]:
    """
    验证合同完整性。按合同类型使用对应关键词集合，并检测残留占位符。

    Args:
        contract_path: 合同文件路径
        contract_type: 合同类型（employment/offer_letter/rental/nda/service/generic）
                       不传则从文件名自动推断
    Returns:
        {'valid': bool, 'warnings': [...], 'errors': [...], 'contract_type': str}
    """
    result: Dict[str, Any] = {
        'valid': True,
        'warnings': [],
        'errors': [],
        'contract_type': contract_type or 'unknown',
    }

    path = Path(contract_path)
    if not path.exists():
        result['valid'] = False
        result['errors'].append(f"文件不存在：{contract_path}")
        return result

    if not DOCX_AVAILABLE:
        result['warnings'].append("python-docx 未安装，跳过内容验证")
        return result

    detected_type = contract_type or _detect_contract_type(contract_path)
    result['contract_type'] = detected_type
    keywords = CONTRACT_KEYWORDS.get(detected_type, CONTRACT_KEYWORDS['generic'])

    doc = Document(contract_path)
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    for keyword in keywords:
        if keyword not in full_text:
            result['warnings'].append(f"未找到关键词：「{keyword}」")

    # 残留占位符检测
    placeholder_pattern = re.compile(r'\[([a-zA-Z_][a-zA-Z0-9_]*)\]')
    unreplaced: set = set()
    for para in doc.paragraphs:
        unreplaced.update(placeholder_pattern.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                unreplaced.update(placeholder_pattern.findall(cell.text))
    for field in sorted(unreplaced):
        result['warnings'].append(f"存在未替换的占位符：[{field}]")

    if result['warnings']:
        result['valid'] = False

    return result


# ---------------------------------------------------------------------------
# list_templates（1.4.0 新增）
# ---------------------------------------------------------------------------

def list_templates(template_dir: Optional[str] = None) -> List[Dict[str, str]]:
    """
    列出所有可用的 DOCX 模板文件。
    返回包含文件名、类型、大小的字典列表。
    """
    tdir = Path(template_dir) if template_dir else Path(__file__).parent / 'templates'
    if not tdir.exists():
        return []

    templates = []
    type_labels = {
        'employment': '雇佣合同',
        'offer_letter': '录用通知书',
        'rental': '租房合同',
        'nda': '保密协议 (NDA)',
        'service': '服务合同',
        'generic': '通用合同',
    }
    for f in sorted(tdir.glob('*.docx')):
        ctype = _detect_contract_type(f.name)
        size_kb = f.stat().st_size / 1024
        templates.append({
            'filename': f.name,
            'type': ctype,
            'type_label': type_labels.get(ctype, '通用合同'),
            'size': f'{size_kb:.1f} KB',
        })
    return templates


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description='Universal Contract Generator v' + __version__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  %(prog)s list-templates
  %(prog)s generate -t employment_contract.docx -v vars.json -o contract.docx
  %(prog)s batch -t offer_letter.docx -c new_hires.csv -o offers/
  %(prog)s batch -t offer_letter.docx -e new_hires.xlsx -o offers/
  %(prog)s validate contract.docx
  %(prog)s validate contract.docx --type rental
        '''
    )
    parser.add_argument('--version', '-V', action='version', version=f'%(prog)s {__version__}')
    subparsers = parser.add_subparsers(dest='command', help='子命令')

    # list-templates（1.4.0 新增）
    ltp = subparsers.add_parser('list-templates', help='列出所有可用模板')
    ltp.add_argument('--template-dir', '-d', help='模板目录（默认为 templates/）')

    # generate
    gp = subparsers.add_parser('generate', help='生成单个合同')
    gp.add_argument('--template', '-t', required=True, help='模板文件名')
    gp.add_argument('--vars', '-v', required=True, help='变量 JSON 文件')
    gp.add_argument('--output', '-o', required=True, help='输出文件路径')
    gp.add_argument('--template-dir', '-d', help='模板目录')

    # batch
    bp = subparsers.add_parser('batch', help='批量生成合同')
    bp.add_argument('--template', '-t', required=True, help='模板文件名')
    bp.add_argument('--template-dir', '-d', help='模板目录')
    bp.add_argument('--output-dir', '-o', default='contracts/', help='输出目录（默认：contracts/）')
    bp.add_argument('--name-field', '-n', default='name', help='用于文件名的字段名')
    src = bp.add_mutually_exclusive_group(required=True)
    src.add_argument('--csv', '-c', help='CSV 数据文件')
    src.add_argument('--excel', '-e', help='Excel (.xlsx/.xls) 数据文件')
    src.add_argument('--json', '-j', help='JSON 数据文件')

    # validate
    vp = subparsers.add_parser('validate', help='验证合同完整性')
    vp.add_argument('contract', help='合同文件路径')
    vp.add_argument('--type', dest='contract_type',
                    choices=list(CONTRACT_KEYWORDS.keys()),
                    help='合同类型（不指定则从文件名推断）')

    args = parser.parse_args()

    # --- list-templates ---
    if args.command == 'list-templates':
        templates = list_templates(getattr(args, 'template_dir', None))
        if not templates:
            print("未找到任何模板文件。")
            return 0
        print(f"\n可用模板（共 {len(templates)} 个）：")
        print(f"{'文件名':<35} {'类型':<20} {'大小'}")
        print('-' * 65)
        for t in templates:
            print(f"{t['filename']:<35} {t['type_label']:<20} {t['size']}")
        print()
        return 0

    # --- generate ---
    elif args.command == 'generate':
        try:
            gen = ContractGenerator(args.template_dir)
            gen.load_template(args.template)
            with open(args.vars, 'r', encoding='utf-8') as f:
                variables = json.load(f)
            if not isinstance(variables, dict):
                print("错误：变量文件必须是 JSON 对象格式")
                return 1
            gen.set_variables(variables)
            gen.save(args.output)
            return 0
        except FileNotFoundError as e:
            print(f"错误：{e}")
            return 1
        except json.JSONDecodeError as e:
            print(f"错误：JSON 格式不正确 - {e}")
            return 1
        except Exception as e:
            print(f"错误：生成失败 - {e}")
            return 1

    # --- batch ---
    elif args.command == 'batch':
        try:
            batch = BatchGenerator(args.template, args.template_dir)
            if args.csv:
                batch.load_csv(args.csv)
            elif args.excel:
                batch.load_excel(args.excel)
            elif args.json:
                batch.load_json(args.json)
            result = batch.generate_all(args.output_dir, args.name_field)
            return 0 if result['failed'] == 0 else 1
        except Exception as e:
            print(f"错误：批量生成失败 - {e}")
            return 1

    # --- validate ---
    elif args.command == 'validate':
        result = validate_contract(args.contract, getattr(args, 'contract_type', None))
        ctype = result.get('contract_type', 'unknown')
        print(f"\n合同文件：{args.contract}")
        print(f"合同类型：{ctype}")

        if result['valid'] and not result['warnings'] and not result['errors']:
            print("✓ 合同验证通过，未发现问题。")
            return 0

        print("⚠ 合同验证完成（存在问题）：")
        if result['errors']:
            print("\n  错误：")
            for e in result['errors']:
                print(f"    ✗ {e}")
        if result['warnings']:
            print("\n  警告：")
            for w in result['warnings']:
                print(f"    ! {w}")
        print()
        return 1 if result['errors'] else 0

    else:
        parser.print_help()
        return 0


if __name__ == '__main__':
    sys.exit(main())
