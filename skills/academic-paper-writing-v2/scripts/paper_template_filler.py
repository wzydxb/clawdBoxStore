#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
论文模板填充工具

功能：根据结构化数据填充论文模板，生成标准格式文档
支持模板：期刊论文、学位论文、会议论文

使用方法：
    python paper_template_filler.py --template journal --data paper_data.json --output paper.docx

输入格式：JSON格式的论文结构化数据
"""

import argparse
import json
import os
import sys
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 缺少 python-docx 库")
    print("请安装: pip install python-docx==0.8.11")
    sys.exit(1)


class PaperTemplateFiller:
    """论文模板填充器"""

    # 标准论文结构模板
    TEMPLATES = {
        "journal": {
            "name": "期刊论文",
            "sections": [
                "摘要",
                "关键词",
                "引言",        # 含文献综述，期刊论文通常不独立成章
                "研究方法",
                "研究结果",
                "讨论",
                "结论",
                "参考文献"
            ]
        },
        "thesis": {
            "name": "学位论文",
            "sections": [
                "摘要",
                "Abstract",
                "目录",
                "第1章 绪论",
                "第2章 理论基础与文献综述",
                "第3章 研究设计与方法",
                "第4章 研究结果与分析",
                "第5章 讨论",
                "第6章 结论与展望",
                "参考文献",
                "附录",
                "致谢"
            ]
        },
        "conference": {
            "name": "会议论文",
            "sections": [
                "摘要",
                "引言",
                "相关工作",
                "方法",
                "实验",
                "结论",
                "参考文献"
            ]
        }
    }

    def __init__(self, template_type: str = "journal"):
        """
        初始化填充器

        Args:
            template_type: 模板类型 (journal/thesis/conference)
        """
        self.template_type = template_type.lower()

        if self.template_type not in self.TEMPLATES:
            raise ValueError(f"不支持的模板类型: {template_type}。支持类型: {', '.join(self.TEMPLATES.keys())}")

        self.template_info = self.TEMPLATES[self.template_type]
        self.doc = Document()

        # 设置默认字体
        self._setup_default_font()

    @staticmethod
    def _set_east_asia_font(element, font_name: str) -> None:
        """
        安全设置东亚字体（中文字体）。

        直接访问 element.rPr.rFonts 在某些情况下会因 rPr 为 None 而崩溃。
        此方法先确保 rPr 节点存在，再设置 eastAsia 字体属性。
        """
        from lxml import etree
        # 确保 rPr 节点存在
        rPr = element.get_or_add_rPr() if hasattr(element, 'get_or_add_rPr') else element.rPr
        if rPr is None:
            # 手动创建 rPr 节点
            rPr = etree.SubElement(element, qn('w:rPr'))
        # 确保 rFonts 子节点存在
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)

    def _setup_default_font(self):
        """设置默认字体（安全版本，避免 rPr 为 None 崩溃）"""
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        # 安全设置中文字体
        self._set_east_asia_font(style.element, '宋体')

    def _add_heading(self, text: str, level: int = 1):
        """
        添加标题

        Args:
            text: 标题文本
            level: 标题级别 (1-4)
        """
        heading = self.doc.add_heading(text, level=level)

        for run in heading.runs:
            run.font.name = 'Times New Roman'
            self._set_east_asia_font(run._element, '黑体')

    def _add_paragraph(self, text: str, bold: bool = False, indent: bool = True):
        """
        添加段落

        Args:
            text: 段落文本
            bold: 是否加粗
            indent: 是否首行缩进
        """
        para = self.doc.add_paragraph()

        if indent:
            para.paragraph_format.first_line_indent = Cm(0.74)

        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        self._set_east_asia_font(run._element, '宋体')

        if bold:
            run.bold = True

    def _add_section(self, section_title: str, content: str):
        """
        添加章节

        Args:
            section_title: 章节标题
            content: 章节内容
        """
        self._add_heading(section_title, level=2)

        if content:
            # 按段落分割内容
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    self._add_paragraph(para_text.strip())

    def fill_title_page(self, data: Dict):
        """
        填充封面页

        Args:
            data: 论文数据字典
        """
        # 添加标题
        title = data.get("title", "论文标题")
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'
        self._set_east_asia_font(title_run._element, '黑体')

        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 添加作者信息
        authors = data.get("authors", [])
        if authors:
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(", ".join(authors))
            author_run.font.size = Pt(14)

        # 添加机构信息
        institution = data.get("institution", "")
        if institution:
            inst_para = self.doc.add_paragraph()
            inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_run = inst_para.add_run(institution)
            inst_run.font.size = Pt(12)

        # 添加日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = data.get("date", datetime.now().strftime("%Y年%m月"))
        date_run = date_para.add_run(date_str)
        date_run.font.size = Pt(12)

        # 添加分页符
        self.doc.add_page_break()

    def fill_abstract(self, data: Dict):
        """
        填充摘要部分

        Args:
            data: 论文数据字典
        """
        self._add_heading("摘要", level=1)

        abstract = data.get("abstract", "")
        if abstract:
            self._add_paragraph(abstract, indent=True)

        # 添加关键词
        keywords = data.get("keywords", [])
        if keywords:
            kw_para = self.doc.add_paragraph()
            kw_run1 = kw_para.add_run("关键词：")
            kw_run1.bold = True
            kw_run1.font.name = 'Times New Roman'
            self._set_east_asia_font(kw_run1._element, '黑体')

            kw_run2 = kw_para.add_run("; ".join(keywords))
            kw_run2.font.name = 'Times New Roman'
            self._set_east_asia_font(kw_run2._element, '宋体')

        self.doc.add_page_break()

    def fill_sections(self, data: Dict):
        """
        填充正文章节

        Args:
            data: 论文数据字典
        """
        sections = data.get("sections", {})

        for section_title in self.template_info["sections"]:
            # 跳过已处理的特殊章节
            if section_title in ["摘要", "关键词", "Abstract"]:
                continue

            # 获取章节内容
            content = sections.get(section_title, "")

            if content:
                self._add_section(section_title, content)
            else:
                # 空章节也添加标题
                self._add_heading(section_title, level=2)
                self._add_paragraph("[待补充内容]", indent=True)

    def fill_references(self, data: Dict):
        """
        填充参考文献

        Args:
            data: 论文数据字典
        """
        self._add_heading("参考文献", level=1)

        references = data.get("references", [])

        if references:
            for ref in references:
                ref_para = self.doc.add_paragraph()
                ref_para.paragraph_format.first_line_indent = Cm(-0.74)  # 悬挂缩进
                ref_para.paragraph_format.left_indent = Cm(0.74)

                ref_run = ref_para.add_run(ref)
                ref_run.font.name = 'Times New Roman'
                ref_run.font.size = Pt(10.5)
                self._set_east_asia_font(ref_run._element, '宋体')
        else:
            self._add_paragraph("[待补充参考文献]", indent=False)

    def fill_document(self, data: Dict) -> Document:
        """
        填充完整文档

        Args:
            data: 论文数据字典

        Returns:
            填充后的Document对象
        """
        # 填充封面
        self.fill_title_page(data)

        # 填充摘要
        self.fill_abstract(data)

        # 填充正文
        self.fill_sections(data)

        # 填充参考文献
        self.fill_references(data)

        return self.doc

    def save(self, output_path: str):
        """
        保存文档

        Args:
            output_path: 输出文件路径
        """
        self.doc.save(output_path)


def validate_input_data(data: Dict) -> bool:
    """
    验证输入数据格式

    Args:
        data: 输入数据字典

    Returns:
        验证是否通过
    """
    required_fields = ["title"]

    for field in required_fields:
        if field not in data:
            print(f"错误: 输入数据缺少 '{field}' 字段")
            return False

    return True


def create_sample_data() -> Dict:
    """创建示例数据"""
    return {
        "title": "基于深度学习的图像分类方法研究",
        "authors": ["张三", "李四", "王五"],
        "institution": "XX大学计算机科学与技术学院",
        "date": "2024年1月",
        "abstract": "本研究提出了一种基于深度学习的图像分类方法，通过改进卷积神经网络结构，在多个公开数据集上取得了优异的分类性能。实验结果表明，该方法相比传统方法准确率提升了5.3%，具有重要的理论意义和应用价值。",
        "keywords": ["深度学习", "图像分类", "卷积神经网络", "特征提取"],
        "sections": {
            "引言": "随着人工智能技术的快速发展，深度学习在计算机视觉领域取得了显著成果...\n\n本文的主要贡献包括：（1）提出了一种新的网络结构；（2）在多个数据集上验证了方法的有效性。",
            "文献综述": "近年来，深度学习在图像分类任务中得到了广泛应用。Krizhevsky等人提出的AlexNet模型开启了深度学习在计算机视觉中的新篇章...",
            "研究方法": "本研究采用对比实验方法，在ImageNet、CIFAR-10等标准数据集上进行实验...",
            "研究结果": "实验结果显示，本文提出的方法在准确率、召回率等指标上均优于对比方法...",
            "讨论": "本研究的结果表明，改进的网络结构能够更有效地提取图像特征...",
            "结论": "本文提出了一种基于深度学习的图像分类方法，通过实验验证了其有效性。未来的研究方向包括..."
        },
        "references": [
            "[1] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[J]. Advances in neural information processing systems, 2012, 25: 1097-1105.",
            "[2] Simonyan K, Zisserman A. Very deep convolutional networks for large-scale image recognition[J]. arXiv preprint arXiv:1409.1556, 2014.",
            "[3] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE conference on computer vision and pattern recognition. 2016: 770-778."
        ]
    }


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="论文模板填充工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python paper_template_filler.py --template journal --data paper_data.json --output paper.docx
    python paper_template_filler.py --template thesis --data thesis_data.json --output thesis.docx
    python paper_template_filler.py --sample  # 生成示例文档
        """
    )

    parser.add_argument(
        "--template", "-t",
        required=False,
        default="journal",
        choices=["journal", "thesis", "conference"],
        help="论文模板类型 (默认: journal)"
    )

    parser.add_argument(
        "--data", "-d",
        required=False,
        help="输入论文数据文件路径（JSON格式）"
    )

    parser.add_argument(
        "--output", "-o",
        required=False,
        default="output.docx",
        help="输出文档路径 (默认: output.docx)"
    )

    parser.add_argument(
        "--sample",
        action="store_true",
        help="生成示例文档"
    )

    args = parser.parse_args()

    # 生成示例文档
    if args.sample:
        print("正在生成示例文档...")
        data = create_sample_data()
        filler = PaperTemplateFiller("journal")
        filler.fill_document(data)
        filler.save(args.output)
        print(f"示例文档已生成: {args.output}")
        return

    # 检查是否提供了数据文件
    if not args.data:
        print("错误: 请使用 --data 参数提供数据文件，或使用 --sample 生成示例文档")
        parser.print_help()
        sys.exit(1)

    # 检查输入文件是否存在
    if not os.path.exists(args.data):
        print(f"错误: 输入文件不存在: {args.data}")
        sys.exit(1)

    # 读取输入数据
    try:
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"错误: JSON解析失败: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"错误: 读取文件失败: {e}")
        sys.exit(1)

    # 验证数据格式
    if not validate_input_data(data):
        sys.exit(1)

    # 填充模板
    try:
        filler = PaperTemplateFiller(args.template)
        filler.fill_document(data)
        filler.save(args.output)

        print(f"成功! 论文文档已生成: {args.output}")
        print(f"  - 模板类型: {filler.template_info['name']}")
        print(f"  - 标题: {data.get('title', '未命名')}")

    except Exception as e:
        print(f"错误: 文档生成失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
